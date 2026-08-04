from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "proposal" / "DeepAlign-Bench_研究Proposal.md"
FIG = ROOT / "proposal_assets" / "DeepAlign-Bench_主图.png"
OUT = ROOT / "deliverables" / "DeepAlign-Bench_正式研究Proposal.docx"

# The formal proposal is the default. Communication variants override these
# module-level values from build_readable_variants.py while reusing the same
# deterministic layout and Markdown parser.
COVER_KICKER = "RESEARCH PROPOSAL"
COVER_TITLE = "DeepAlign-Bench"
COVER_SUBTITLE = "长程 Deep Research 智能体个性化最终交付物评测"
COVER_MODE = "Benchmark · Evaluation · Human-Centered Agents"
DOC_VERSION = "v0.26 · 组内讨论稿"
DOC_DATE = "2026 年 8 月 4 日"
RESEARCH_LINE = "Counterfactual Effect · 可审计构造 · 分级压力 · 跨环境运行"
CORE_CLAIM = "固定任务与证据，只改变目标用户；用 matched/swapped 与三类契约判断输出是否按正确方向变化，再比较不同 agent 在 S0–S3 压力下的表现。"
CONTENTS_ITEMS = [
    "研究概要与可证伪假设", "关键文献精读与设计启示", "Evaluation Atlas 与双轴 taxonomy",
    "Benchmark 数据结构与构建流程", "Rubric、Metrics 与 Judge", "实验矩阵与平台实现",
    "严格审稿风险与防守", "里程碑、论文结构与最小可行版本", "论文图表蓝图",
    "参考文献",
]
READING_NOTE = "阅读提示：主图给出整体逻辑；第 7–8 节是本 proposal 的测量学核心；第 11 节按顶会审稿视角集中列出可预见攻击与防守。"
FIGURE_TRIGGER = "2. 关键文献"
FIGURE_TITLE = "总体框架：从受控用户信号到反事实评估"
FIGURE_CAPTION = "图 1  DeepAlign-Bench 主流程。主榜先过共同质量门槛，再检查跨用户对角优势与跨 cue 稳健性；JudgeBench 独立验证自动评委。"
RUNNING_HEADER = "DEEPALIGN-BENCH  ·  RESEARCH PROPOSAL"
STYLE_PRESET = "narrative_proposal"
INCLUDE_CONTENTS = True

BLUE = "2E74B5"
DARK = "163A63"
MUTED = "5B6B7A"
LIGHT = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "B68026"
RED = "9B3A2A"
FONT = "Calibri"
CN_FONT = "PingFang SC"
HEADING_CN_FONT = "Hiragino Sans GB"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=None, bold=None, color=None, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, field):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_font(run, size=9, color=MUTED)


def configure_section(section, landscape=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        margin = Inches(0.65)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        if STYLE_PRESET == "formal_condensed":
            margin = Inches(0.75)
        elif STYLE_PRESET == "compact_reference_guide":
            margin = Inches(0.75)
        else:
            margin = Inches(1)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    p = header.paragraphs[0]
    p.text = RUNNING_HEADER
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        set_font(r, size=8.5, bold=True, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    label = p.add_run(f"{DOC_DATE.replace(' 年 ', '-').replace(' 月 ', '-').replace(' 日', '')}   ·   ")
    set_font(label, size=9, color=MUTED)
    add_field(p, "PAGE")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    heading_tokens = (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    )
    if STYLE_PRESET == "compact_reference_guide":
        normal.font.size = Pt(9.7)
        normal.paragraph_format.space_after = Pt(2.5)
        normal.paragraph_format.line_spacing = 1.05
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_tokens = (
            ("Heading 1", 15.5, BLUE, 14, 7),
            ("Heading 2", 12.5, BLUE, 10, 5),
            ("Heading 3", 11.5, DARK, 7, 4),
        )
    elif STYLE_PRESET == "formal_condensed":
        # Named override for the <=10-page academic proposal edition. The
        # hierarchy remains formal; only paragraph rhythm is tightened.
        normal.font.size = Pt(10)
        normal.paragraph_format.space_after = Pt(3)
        normal.paragraph_format.line_spacing = 1.08
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_tokens = (
            ("Heading 1", 15, BLUE, 12, 6),
            ("Heading 2", 12.5, BLUE, 8, 4),
            ("Heading 3", 11.5, DARK, 6, 3),
        )
    for name, size, color, before, after in heading_tokens:
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        # PingFang's synthetic bold is rendered as solid blocks for a few
        # high-stroke CJK glyphs by LibreOffice. Hiragino supplies a native
        # bold face and keeps headings legible in both DOCX and exported PDF.
        st._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_CN_FONT)
        st.font.size = Pt(size)
        # Size and colour carry the hierarchy. Avoid synthetic CJK bold: in
        # LibreOffice/PDF export dense glyphs such as 量 and 盖 can fill in.
        st.font.bold = False
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
        compact_lists = STYLE_PRESET in {"formal_condensed", "compact_reference_guide"}
        if STYLE_PRESET == "formal_condensed":
            st.font.size = Pt(9.7)
        else:
            st.font.size = Pt(10 if compact_lists else 11)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.194)
        st.paragraph_format.space_after = Pt(2.5 if STYLE_PRESET == "formal_condensed" else (3 if compact_lists else 4))
        st.paragraph_format.line_spacing = 1.05 if STYLE_PRESET == "formal_condensed" else (1.08 if compact_lists else 1.208)


def add_hyperlink(paragraph, label, url, size=None, color=BLUE):
    """Append a clickable external link while preserving the document type system."""
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), FONT)
    run_fonts.set(qn("w:hAnsi"), FONT)
    run_fonts.set(qn("w:eastAsia"), CN_FONT)
    run_props.append(run_fonts)
    link_color = OxmlElement("w:color")
    link_color.set(qn("w:val"), color or BLUE)
    run_props.append(link_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)
    if size is not None:
        half_points = str(int(round(size * 2)))
        font_size = OxmlElement("w:sz")
        font_size.set(qn("w:val"), half_points)
        run_props.append(font_size)
        font_size_cs = OxmlElement("w:szCs")
        font_size_cs.set(qn("w:val"), half_points)
        run_props.append(font_size_cs)
    run.append(run_props)
    node = OxmlElement("w:t")
    if label.startswith(" ") or label.endswith(" "):
        node.set(qn("xml:space"), "preserve")
    node.text = label
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text, size=None, color=None):
    # Small Markdown subset: bold, emphasis, inline code, Markdown links and bare URLs.
    pos = 0
    pattern = re.compile(
        r"(\[\[\d+\]\]\(https?://[^)\s]+\)|\[[^\]]+\]\(https?://[^)\s]+\)|https?://[^\s]+|\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`.+?`)"
    )
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            set_font(r, size=size, color=color)
        token = m.group(0)
        citation_link = re.fullmatch(r"\[\[(\d+)\]\]\((https?://[^)\s]+)\)", token)
        markdown_link = re.fullmatch(r"\[([^\]]+)\]\((https?://[^)\s]+)\)", token)
        if citation_link:
            add_hyperlink(paragraph, f"[{citation_link.group(1)}]", citation_link.group(2), size=size)
        elif markdown_link:
            add_hyperlink(paragraph, markdown_link.group(1), markdown_link.group(2), size=size)
        elif token.startswith("http://") or token.startswith("https://"):
            trailing = ""
            while token and token[-1] in ".,;:!?，。；：！？":
                trailing = token[-1] + trailing
                token = token[:-1]
            add_hyperlink(paragraph, token, token, size=size)
            if trailing:
                r = paragraph.add_run(trailing)
                set_font(r, size=size, color=color)
        elif token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            set_font(r, size=size, bold=True, color=color)
        elif token.startswith("*"):
            r = paragraph.add_run(token[1:-1])
            set_font(r, size=size, italic=True, color=color)
        else:
            r = paragraph.add_run(token[1:-1])
            set_font(r, size=size or 10, color="6A3E00", name="Courier New")
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_font(r, size=size, color=color)


def paragraph_border_bottom(paragraph, color=BLUE, size="12", space="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(COVER_KICKER)
    set_font(r, size=11, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(COVER_TITLE)
    set_font(r, size=30, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(COVER_SUBTITLE)
    set_font(r, size=17, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run(COVER_MODE)
    set_font(r, size=10.5, bold=True, color=MUTED)
    paragraph_border_bottom(p, color="7A9BB7", size="8", space="8")
    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, [4680, 4680])
    values = [
        ("文档版本", DOC_VERSION),
        ("日期", DOC_DATE),
        ("研究主线", RESEARCH_LINE),
    ]
    for i, (a, b) in enumerate(values):
        for j, value in enumerate((a, b)):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(value)
            set_font(r, size=10.5, bold=(j == 0), color=DARK if j == 0 else MUTED)
            if i % 2 == 0:
                set_cell_shading(cell, LIGHT)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("核心命题")
    set_font(r, size=11, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(CORE_CLAIM)
    set_font(r, size=12, bold=True, color=DARK)
    doc.add_page_break()


def add_contents(doc):
    p = doc.add_paragraph("内容导航", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    items = CONTENTS_ITEMS
    num_id = new_numbering_instance(doc)
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        apply_numbering(p, num_id)
        p.paragraph_format.space_after = Pt(5)
        add_inline(p, item)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(READING_NOTE)
    set_font(r, size=10.5, color=MUTED, italic=True)
    doc.add_page_break()


def add_table(doc, rows):
    ncols = max(len(r) for r in rows)
    # Content-specific patterns; all widths sum to 9360 DXA.
    if ncols == 4:
        widths = [1500, 3000, 2520, 2340]
    elif ncols == 3:
        widths = [1700, 3000, 4660]
    else:
        widths = [9360 // ncols] * ncols
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    table_font_size = 8.2 if STYLE_PRESET == "compact_reference_guide" else 9.2
    table_spacing = 0 if STYLE_PRESET == "compact_reference_guide" else 2
    table_line_spacing = 1.0 if STYLE_PRESET == "compact_reference_guide" else 1.08
    for i, row in enumerate(rows):
        for j in range(ncols):
            text = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(table_spacing)
            p.paragraph_format.line_spacing = table_line_spacing
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (i == 0 or j == 0) else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, text, size=table_font_size, color=DARK if i == 0 else None)
            if i == 0:
                set_cell_shading(cell, LIGHT)
                for run in p.runs:
                    run.bold = True
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    return table


def new_numbering_instance(doc, abstract_num_id=0):
    numbering = doc.part.numbering_part.element
    existing = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = max(existing or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.194)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208


def add_figure_section(doc):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(sec, landscape=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(FIGURE_TITLE)
    set_font(r, size=17, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    picture = run.add_picture(str(FIG), width=Inches(9.55))
    picture._inline.docPr.set(
        "descr",
        "DeepAlign-Bench 研究流程：从五平面元数据、反事实任务族和跨 agent 运行，到 rubric compiler、分层 judge 与人评校准。",
    )
    picture._inline.docPr.set("title", "DeepAlign-Bench 研究设计")
    p.paragraph_format.space_after = Pt(5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(FIGURE_CAPTION)
    set_font(r, size=9.5, color=MUTED, italic=True)
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(sec2, landscape=False)


def build(md_path=MD, out_path=OUT):
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], landscape=False)
    add_cover(doc)
    if INCLUDE_CONTENTS:
        add_contents(doc)

    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    # Skip the Markdown title and metadata. Formal variants may open with either
    # a research overview or a conventional abstract.
    start = next(
        i
        for i, line in enumerate(lines)
        if line.strip() in {"## 研究概要", "## 摘要"}
    )
    lines = lines[start:]
    paragraph_buf = []
    in_code = False
    code_buf = []
    figure_added = False
    current_num_id = None

    def flush_paragraph():
        nonlocal paragraph_buf
        if paragraph_buf:
            text = " ".join(x.strip() for x in paragraph_buf).strip()
            if text:
                p = doc.add_paragraph()
                add_inline(p, text)
            paragraph_buf = []

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            if not in_code:
                in_code = True
                code_buf = []
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.15)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.0
                p_pr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "EEF2F5")
                p_pr.append(shd)
                r = p.add_run("\n".join(code_buf))
                set_font(r, size=9.2, color=DARK, name="Courier New")
                in_code = False
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        if not stripped:
            flush_paragraph()
            current_num_id = None
            i += 1
            continue
        if stripped == "---":
            flush_paragraph()
            i += 1
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            title = stripped[3:]
            if title.startswith(FIGURE_TRIGGER) and not figure_added:
                # Figure follows the research questions and precedes literature review.
                add_figure_section(doc)
                figure_added = True
            doc.add_paragraph(title, style="Heading 1")
            current_num_id = None
            i += 1
            continue
        if stripped.startswith("### "):
            flush_paragraph()
            doc.add_paragraph(stripped[4:], style="Heading 2")
            current_num_id = None
            i += 1
            continue
        if stripped.startswith("#### "):
            flush_paragraph()
            doc.add_paragraph(stripped[5:], style="Heading 3")
            current_num_id = None
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[i + 1]):
            flush_paragraph()
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                parts = [x.strip() for x in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", p.replace(" ", "")) for p in parts):
                    rows.append(parts)
                i += 1
            add_table(doc, rows)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        if re.match(r"^\[\d+\]\s", stripped):
            flush_paragraph()
            p = doc.add_paragraph()
            if STYLE_PRESET == "formal_condensed":
                reference_size = 7.6
            elif STYLE_PRESET == "compact_reference_guide":
                reference_size = 7.4
            else:
                reference_size = 8.4
            add_inline(p, stripped, size=reference_size)
            current_num_id = None
            i += 1
            continue
        m_bullet = re.match(r"^\s*-\s+(.*)", line)
        m_num = re.match(r"^\s*\d+\.\s+(.*)", line)
        if m_bullet or m_num:
            flush_paragraph()
            if m_bullet:
                p = doc.add_paragraph(style="List Bullet")
                current_num_id = None
            else:
                if current_num_id is None:
                    current_num_id = new_numbering_instance(doc)
                p = doc.add_paragraph()
                apply_numbering(p, current_num_id)
            add_inline(p, (m_bullet or m_num).group(1))
            i += 1
            continue
        paragraph_buf.append(stripped)
        i += 1
    flush_paragraph()

    # Keep references compact and ensure final paragraph doesn't strand.
    for p in doc.paragraphs:
        if p.text.startswith("[") and re.match(r"^\[\d+\]", p.text):
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            if STYLE_PRESET == "formal_condensed":
                ref_size, ref_space, ref_line = 7.6, 0, 0.95
            elif STYLE_PRESET == "compact_reference_guide":
                ref_size, ref_space, ref_line = 7.4, 0, 0.9
            else:
                ref_size, ref_space, ref_line = 8.4, 0, 0.92
            p.paragraph_format.space_after = Pt(ref_space)
            p.paragraph_format.line_spacing = ref_line
            for r in p.runs:
                set_font(r, size=ref_size)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    build()
